from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.shortcuts import render, redirect
from django.contrib import messages
from django.db import transaction
from django.core.exceptions import ObjectDoesNotExist

from accounts.models import (
    StudentAssessment,
    User,
    Administrator,
    Staff,
    Student,
    Session,
    Term,
    ClassLevel,
    Subject,
    Fee,
    Payment,
)
from accounts.forms import AddStudentForm, EditStudentForm, FeeForm
from accounts.utils import current_term, generate_school_id

import docx
import json
from io import StringIO


def home(request, **kwargs):
    user = User.objects.get(school_id=kwargs.get("user_school_id"))

    # data used in every view
    request.session["user_school_id"] = user.school_id
    request.session["user_first_name"] = user.first_name
    request.session["user_last_name"] = user.last_name
    request.session["user_other_names"] = user.other_names

    all_student_count = Student.objects.all().count()
    subject_count = Subject.objects.all().count()
    class_level_count = ClassLevel.objects.all().count()
    staff_count = Staff.objects.all().count()

    # Total Subjects and students in Each Course
    class_level_all = ClassLevel.objects.all()
    class_level_name_list = []
    subject_count_list = []
    student_count_list_in_class_level = []

    for class_level in class_level_all:
        subjects = Subject.objects.filter(class_level_id=class_level.id).count()
        students = Student.objects.filter(class_level_id=class_level.id).count()
        class_level_name_list.append(class_level.class_level_name)
        subject_count_list.append(subjects)
        student_count_list_in_class_level.append(students)

    subject_all = Subject.objects.all()
    subject_list = []
    student_count_list_in_subject = []
    for subject in subject_all:
        class_level = ClassLevel.objects.get(id=subject.class_level.id)
        student_count = Student.objects.filter(class_level_id=class_level.id).count()
        subject_list.append(subject.subject_name)
        student_count_list_in_subject.append(student_count)

    # For Saffs
    staff_name_list = []

    staffs = Staff.objects.all()
    for staff in staffs:
        staff_name_list.append(staff.user.first_name)

    # For Students
    student_name_list = []

    students = Student.objects.all()
    for student in students:
        student_name_list.append(student.user.first_name)

    context = {
        "user_school_id": request.session.get("user_school_id"),
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "all_student_count": all_student_count,
        "subject_count": subject_count,
        "class_level_count": class_level_count,
        "staff_count": staff_count,
        "class_level_name_list": class_level_name_list,
        "subject_count_list": subject_count_list,
        "student_count_list_in_class_level": student_count_list_in_class_level,
        "subject_list": subject_list,
        "student_count_list_in_subject": student_count_list_in_subject,
        "staff_name_list": staff_name_list,
        "student_name_list": student_name_list,
        "courses": class_level_all,
    }

    return render(request, "admin_templates/home.html", context)


# Administrator profile


def profile(request, user_school_id):
    user = User.objects.get(school_id=user_school_id)

    context = {
        "user": user,
        "user_school_id": request.session.get("user_school_id"),
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "gender_data": User.gender_data,
    }
    return render(request, "admin_templates/admin_profile.html", context)


def edit_admin_profile(request, user_school_id):

    if request.method != "POST":
        messages.error(request, "Invalid Method!")
        return redirect("/" + user_school_id + "profile")

    else:
        first_name = request.POST.get("first_name")
        last_name = request.POST.get("last_name")
        other_names = request.POST.get("other_names")
        password = request.POST.get("password")
        email = request.POST.get("email")
        phone_number = request.POST.get("phone_number")
        school_id = request.POST.get("school_id")
        gender = request.POST.get("gender")
        profile_pic = request.FILES.get("profile_pic")

        try:
            user = User.objects.get(school_id=school_id)
            user.first_name = first_name
            user.last_name = last_name
            user.other_names = other_names
            user.email = email
            user.phone_number = phone_number
            user.gender = gender

            if password != None and password != "":
                user.set_password(password)

            if profile_pic != None or profile_pic != "":
                user.profile_pic = profile_pic

            user.save()
            messages.success(request, "Profile Updated Successfully")

        except:
            messages.error(request, "Failed to Update Profile")

        finally:
            return redirect("/" + user_school_id + "/profile")


# STAFF


def manage_staff(request, user_school_id):
    staff_list = Staff.objects.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "staff_list": staff_list,
    }

    return render(request, "admin_templates/manage_staff_template.html", context)


def add_staff(request, user_school_id):
    gender_data = User.gender_data
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "gender_data": gender_data,
    }

    return render(request, "admin_templates/add_staff_template.html", context)


def add_staff_save(request, user_school_id):
    if request.method != "POST":
        messages.error(request, "Invalid Method ")
        return redirect("/" + user_school_id + "/add_staff")

    else:
        first_name = request.POST.get("first_name")
        last_name = request.POST.get("last_name")
        other_names = request.POST.get("other_names")
        email = request.POST.get("email")
        password = request.POST.get("password")
        phone_number = request.POST.get("phone_number")
        gender = request.POST.get("gender")
        username = generate_school_id()

        try:
            user = User.objects.create_user(
                username=username,
                password=password,
                email=email,
                first_name=first_name,
                last_name=last_name,
                other_names=other_names,
                phone_number=phone_number,
                gender=gender,
                user_type=2,
            )

            messages.success(request, "Staff Added Successfully!")

        except Exception as e:
            print(e)
            messages.error(request, "Failed to Add Staff!")

        finally:
            return redirect("/" + user_school_id + "/add_staff")


def edit_staff(request, user_school_id, staff_school_id):

    staff = User.objects.get(school_id=staff_school_id).staff
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "staff": staff,
        "gender_data": User.gender_data,
    }
    return render(request, "admin_templates/edit_staff_template.html", context)


def edit_staff_save(request, user_school_id):

    if request.method != "POST":
        return HttpResponse("<h2>Method Not Allowed</h2>")

    else:
        staff_school_id = request.POST.get("staff_school_id")
        email = request.POST.get("email")
        first_name = request.POST.get("first_name")
        last_name = request.POST.get("last_name")
        other_names = request.POST.get("other_names")
        phone_number = request.POST.get("phone_number")
        gender = request.POST.get("gender")

        try:
            # INSERTING into User Model
            user = User.objects.get(school_id=staff_school_id)
            user.first_name = first_name
            user.last_name = last_name
            user.email = email
            user.other_names = other_names
            user.phone_number = phone_number
            user.gender = gender
            user.save()

            messages.success(request, "Staff Updated Successfully.")

        except:
            messages.error(request, "Failed to Update Staff.")

        finally:
            return redirect("/" + user_school_id + "/edit_staff/" + staff_school_id)


def delete_staff(request, user_school_id, staff_school_id):
    staff = User.objects.get(school_id=staff_school_id)

    try:
        staff.delete()
        messages.success(request, "Staff Deleted Successfully.")

    except:
        messages.error(request, "Failed to Delete Staff.")

    finally:
        return redirect("/" + user_school_id + "/manage_staff")


# STUDENTS


def manage_student(request, user_school_id):

    students = Student.objects.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "students": students,
    }
    return render(request, "admin_templates/manage_students_template.html", context)


def add_student(request, user_school_id):
    form = AddStudentForm()

    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "form": form,
    }
    return render(request, "admin_templates/add_student_template.html", context)


def add_student_save(request, user_school_id):

    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect("/" + user_school_id + "/add_student")

    else:
        form = AddStudentForm(request.POST, request.FILES)

        try :
            class_level_id = request.POST['class_level_id']
        except KeyError:
            messages.error(request, "Select a Class To Add Student!") 
            return redirect("/" + user_school_id + "/add_student")

        if form.is_valid():
            first_name = form.cleaned_data["first_name"]
            last_name = form.cleaned_data["last_name"]
            other_names = form.cleaned_data["other_names"]
            email = form.cleaned_data["email"]
            password = form.cleaned_data["password"]
            phone_number = form.cleaned_data["phone_number"]
            address = form.cleaned_data["address"]
            dob = form.cleaned_data["dob"]
            class_level_id = (
                form.cleaned_data["class_level_id"]
                if form.cleaned_data["class_level_id"]
                else request.session.get("class_level_id")
            )
            gender = form.cleaned_data["gender"]
            username = generate_school_id()
            profile_pic = form.cleaned_data["profile_pic"]


            try:
                current_term = Term.objects.get(current_term=True)
                
                class_level = ClassLevel.objects.get(id=class_level_id)

                
                
                user = User.objects.create_user(
                username=username,
                password=password,
                email=email,
                first_name=first_name,
                last_name=last_name,
                other_names=other_names,
                user_type=3,
                phone_number=phone_number,
                gender=gender,
                )


                if profile_pic != None or profile_pic != "":
                    user.profile_pic = profile_pic

                user.student.address = address
                user.student.address = dob
                user.student.phone_number = phone_number


                user.student.class_level = class_level

                user.student.term_enrolled = current_term

                user.save()
                messages.success(request, "Student Added Successfully!")

                
                terms_in_current_session = current_term.session_set.all()
                assessment_list = []
                try:
                    old_student_instance = class_level.student_set.all()[0]
                    for term in terms_in_current_session:
                        try:
                            assessments = old_student_instance.studentassessment_set.filter(term=term)
                            assessment_list.append(assessments)
                            
                        except StudentAssessment.DoesNotExist:
                            pass

                    if assessment_list != []:
                        for term_assessments in assessment_list:
                            for assessment in term_assessments:
                                new_assessment_instance = deepcopy(assessment)
                                new_assessment_instance.id = None
                                new_assessment_instance.student = user.student
                                new_assessment_instance.score = 0
                                new_assessment_instance.save()

                except Student.DoesNotExist:
                    pass
            
            except Term.DoesNotExist:
                print('term except')
                messages.error(request, "Select a Current Session and Term To Add Student!") 

            except Exception as E:
                print(E)
                print('general except')
                messages.error(request, "Failed to Add Student!")

            finally:
                return redirect("/" + user_school_id + "/add_student")

        else:
            return redirect("/" + user_school_id + "/add_student")


def edit_student(request, user_school_id, student_school_id):
    # Adding Student ID into Session Variable
    request.session["student_school_id"] = student_school_id

    student_user = User.objects.get(school_id=student_school_id)
    form = EditStudentForm()
    # Filling the form with Data from Database
    form.fields["email"].initial = student_user.email
    form.fields["first_name"].initial = student_user.first_name
    form.fields["last_name"].initial = student_user.last_name
    form.fields["other_names"].initial = student_user.other_names
    form.fields["address"].initial = student_user.student.address
    form.fields["dob"].initial = student_user.student.dob
    form.fields["class_level_id"].initial = student_user.student.class_level.id
    form.fields["gender"].initial = student_user.gender
    form.fields["phone_number"].initial = student_user.phone_number

    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "student": student_user,
        "form": form,
    }
    return render(request, "admin_templates/edit_student_template.html", context)


def edit_student_save(request, user_school_id):
    if request.method != "POST":
        return HttpResponse("Invalid Method!")

    else:
        student_school_id = request.session.get("student_school_id")

        if student_school_id == None:
            return redirect("/" + user_school_id + "/manage_student")

        form = EditStudentForm(request.POST, request.FILES)
        if form.is_valid():
            email = form.cleaned_data["email"]
            first_name = form.cleaned_data["first_name"]
            last_name = form.cleaned_data["last_name"]
            other_names = form.cleaned_data["other_names"]
            address = form.cleaned_data["address"]
            dob = form.cleaned_data["dob"]
            class_level_id = form.cleaned_data["class_level_id"]
            gender = form.cleaned_data["gender"]
            phone_number = form.cleaned_data["phone_number"]
            profile_pic = form.cleaned_data["profile_pic"]

            # Getting Profile Pic first
            # First Check whether the file is selected or not
            # Upload only if file is selected
            # if len(request.FILES) != 0:
            #     profile_pic = request.FILES['profile_pic']
            #     fs = FileSystemStorage()
            #     filename = fs.save(profile_pic.name, profile_pic)
            #     profile_pic_url = fs.url(filename)
            # else:
            #     profile_pic_url = None

            try:
                # First Update into User Model
                user = User.objects.get(school_id=student_school_id)
                user.first_name = first_name
                user.last_name = last_name
                user.other_names = other_names
                user.email = email
                user.phone_number = phone_number
                user.gender = gender

                if profile_pic != None or profile_pic != "":
                    user.profile_pic = profile_pic

                # Then Update Students Table
                student_model = user.student
                student_model.address = address
                student_model.dob = dob

                class_level = ClassLevel.objects.get(id=class_level_id)
                student_model.class_level = class_level

                user.save()

                student_model.save()
                # Delete student_id SESSION after the data is updated
                del request.session["student_school_id"]

                messages.success(request, "Student Updated Successfully!")

            except:
                messages.success(request, "Failed to Update Student.")

            finally:
                return redirect(
                    "/" + user_school_id + "/edit_student/" + student_school_id
                )

        else:
            return redirect("/" + user_school_id + "/edit_student/" + student_school_id)


def delete_student(request, user_school_id, student_school_id):
    student = User.objects.get(school_id=student_school_id)
    try:
        student.delete()
        messages.success(request, "Student Deleted Successfully.")

    except:
        messages.error(request, "Failed to Delete Student.")

    finally:
        return redirect("/" + user_school_id + "/manage_student")


# CLASSES


def manage_class(request, user_school_id):

    class_levels = ClassLevel.objects.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_levels": class_levels,
    }
    return render(request, "admin_templates/manage_class_template.html", context)


def add_class(request, user_school_id):

    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
    }
    return render(request, "admin_templates/add_class_template.html", context)


def add_class_save(request, user_school_id):
    if request.method != "POST":
        messages.error(request, "Invalid Method!")
        return redirect("/" + user_school_id + "/add_class")

    else:
        class_level_name = request.POST.get("class_level_name")
        try:
            course_model = ClassLevel(class_level_name=class_level_name)
            course_model.save()
            messages.success(request, "Class Added Successfully!")

        except:
            messages.error(request, "Failed to Add Class!")

        finally:
            return redirect("/" + user_school_id + "/add_class")


def edit_class(request, user_school_id, class_level_id):
    class_level = ClassLevel.objects.get(id=class_level_id)
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_level": class_level,
    }
    return render(request, "admin_templates/edit_class_template.html", context)


def edit_class_save(request, user_school_id):
    if request.method != "POST":
        HttpResponse("Invalid Method")

    else:
        class_level_id = request.POST.get("class_level_id")
        class_level_name = request.POST.get("class_level_name")

        try:
            class_level = ClassLevel.objects.get(id=class_level_id)
            class_level.class_level_name = class_level_name
            class_level.save()

            messages.success(request, "Class Updated Successfully.")

        except:
            messages.error(request, "Failed to Update Class.")

        finally:
            return redirect("/" + user_school_id + "/edit_class/" + class_level_id)


def manage_class_students(request, user_school_id, class_level_name):
    class_level = ClassLevel.objects.get(class_level_name=class_level_name)
    students = class_level.student_set.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_level": class_level,
        "students": students,
    }

    return render(request, "admin_templates/manage_students_template.html", context)


def manage_class_add_student(request, user_school_id, class_level_name):
    class_level = ClassLevel.objects.get(class_level_name=class_level_name)
    form = AddStudentForm()

    # Filling the form with Data from request.session object when
    # add_student view is accessed via manage_class > manage_students
    form.fields["class_level_id"].initial = class_level.id
    form.fields["class_level_id"].disabled = True

    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "form": form,
        "class_level": class_level,
    }
    return render(request, "admin_templates/add_student_template.html", context)


def manage_class_subjects(request, user_school_id, class_level_name):
    class_level = ClassLevel.objects.get(class_level_name=class_level_name)
    subjects = class_level.subject_set.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_level": class_level,
        "subjects": subjects,
    }

    return render(request, "admin_templates/manage_subjects_template.html", context)


def manage_class_add_subject(request, user_school_id, class_level_name):

    class_levels = ClassLevel.objects.all()
    staff_list = User.objects.filter(user_type="2")
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_levels": class_levels,
        "class_level_name": class_level_name,
        "staff_list": staff_list,
    }
    return render(request, "admin_templates/add_subject_template.html", context)


def delete_class(request, user_school_id, class_level_id):
    class_level = ClassLevel.objects.get(id=class_level_id)
    try:
        class_level.delete()
        messages.success(request, "Class Deleted Successfully.")

    except:
        messages.error(request, "Failed to Delete class.")

    finally:
        return redirect("/" + user_school_id + "/manage_class")


# SUBJECTS


def manage_subject(request, user_school_id):

    subjects = Subject.objects.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "subjects": subjects,
    }
    return render(request, "admin_templates/manage_subjects_template.html", context)


def add_subject(request, user_school_id):

    class_levels = ClassLevel.objects.all()
    staff_list = User.objects.filter(user_type="2")
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_levels": class_levels,
        "staff_list": staff_list,
    }
    return render(request, "admin_templates/add_subject_template.html", context)


def add_subject_save(request, user_school_id):

    if request.method != "POST":
        messages.error(request, "Method Not Allowed!")
        return redirect("/" + user_school_id + "/add_subject")
    else:
        subject_name = request.POST.get("subject")
        class_level_id = request.POST.get("class")
        class_level = ClassLevel.objects.get(id=class_level_id)

        staff_id = request.POST.get("staff")
        staff = User.objects.get(id=staff_id)

        try:
            subject = Subject(
                subject_name=subject_name, class_level=class_level, staff=staff
            )
            subject.save()
            messages.success(request, "Subject Added Successfully!")

        except:
            messages.error(request, "Failed to Add Subject!")

        finally:
            return redirect("/" + user_school_id + "/add_subject")


def edit_subject(request, user_school_id, subject_id):
    subject = Subject.objects.get(id=subject_id)
    class_levels = ClassLevel.objects.all()
    staff_list = User.objects.filter(user_type="2")
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "subject": subject,
        "class_levels": class_levels,
        "staff_list": staff_list,
    }
    return render(request, "admin_templates/edit_subject_template.html", context)


def edit_subject_save(request, user_school_id):
    if request.method != "POST":
        HttpResponse("Invalid Method.")

    else:
        subject_id = request.POST.get("subject_id")
        subject_name = request.POST.get("subject_name")
        class_level_id = request.POST.get("class_level_id")
        staff_id = request.POST.get("staff_id")

        try:
            subject = Subject.objects.get(id=subject_id)
            subject.subject_name = subject_name

            class_level = ClassLevel.objects.get(id=class_level_id)
            subject.class_level = class_level

            staff = User.objects.get(id=staff_id)
            subject.staff = staff

            subject.save()

            messages.success(request, "Subject Updated Successfully.")
            # return HttpResponseRedirect(reverse("edit_subject", kwargs={"subject_id":subject_id}))

        except:
            messages.error(request, "Failed to Update Subject.")
            # return HttpResponseRedirect(reverse("edit_subject", kwargs={"subject_id":subject_id}))

        finally:
            return redirect("/" + user_school_id + "/edit_subject/" + subject_id)


def delete_subject(request, user_school_id, subject_id):
    subject = Subject.objects.get(id=subject_id)
    try:
        subject.delete()
        messages.success(request, "Subject Deleted Successfully.")

    except:
        messages.error(request, "Failed to Delete Subject.")

    finally:
        return redirect("/" + user_school_id + "/manage_subject")


# SESSION


def manage_session(request, user_school_id):

    try:
        sessions = Session.objects.all()
        current_term = Term.objects.get(current_term=True)
        context = {
            "user_first_name": request.session.get("user_first_name"),
            "user_last_name": request.session.get("user_last_name"),
            "user_other_names": request.session.get("user_other_names"),
            "user_school_id": user_school_id,
            "sessions": sessions,
            "current_term": current_term,
        }
        return render(request, "admin_templates/manage_session_template.html", context)

    except Term.DoesNotExist:
        sessions = Session.objects.all()
        context = {
            "user_first_name": request.session.get("user_first_name"),
            "user_last_name": request.session.get("user_last_name"),
            "user_other_names": request.session.get("user_other_names"),
            "user_school_id": user_school_id,
            "sessions": sessions,
        }
        return render(request, "admin_templates/manage_session_template.html", context)


def add_session(request, user_school_id):

    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
    }
    return render(request, "admin_templates/add_session_template.html", context)


def add_session_save(request, user_school_id):
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect("/" + user_school_id + "/manage_session")

    else:
        session_start = request.POST.get("session_start")
        session_end = request.POST.get("session_end")

        try:
            session = Session(session_start=session_start, session_end=session_end)
            session.save()
            term_one = Term(session=session, term=1)
            term_one.save()
            term_two = Term(session=session, term=2)
            term_two.save()
            term_three = Term(session=session, term=3)
            term_three.save()
            messages.success(request, "Session added Successfully!")

        except:
            messages.error(request, "Failed to Add Session")

        finally:
            return redirect("/" + user_school_id + "/manage_session")


def select_session(request, user_school_id):

    sessions = Session.objects.all()
    # try:
    #    current_session = CurrentSession.objects.all()
    # except:
    #    current_session = None
    context = {
        #'current_session': current_session,
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "sessions": sessions,
        "terms": Term.term_data,
    }
    return render(request, "admin_templates/current_session.html", context)


def select_session_save(request, user_school_id):

    if request.method != "POST":
        messages.error(request, "Method Not Allowed!")
        return redirect("/" + user_school_id + "/current_session")
    else:
        session_id = request.POST.get("session")
        term_value = request.POST.get("term")
        session = Session.objects.get(id=session_id)
        term = session.term_set.get(term=int(term_value))

        try:
            prev_current_session = Session.objects.get(current_session=True)
            prev_current_session.current_session = False
            prev_current_session.save()
            prev_current_term = Term.objects.get(current_term=True)
            prev_current_term.current_term = False
            prev_current_term.save()
            session.current_session = True
            session.save()
            term.current_term = True
            term.save()
            messages.success(request, "Current Session updated Successfully!")

        except ObjectDoesNotExist:
            session.current_session = True
            session.save()
            term.current_term = True
            term.save()
            messages.success(request, "Current Session updated Successfully!")

        except:
            messages.error(request, "Failed to Create Current Session!")

        finally:
            return redirect("/" + user_school_id + "/current_session")


def edit_session(request, user_school_id, session_id):
    session = Session.objects.get(id=session_id)
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "session": session,
    }
    return render(request, "admin_templates/edit_session_template.html", context)


def edit_session_save(request, user_school_id):
    if request.method != "POST":
        messages.error(request, "Invalid Method!")
        return redirect("/" + user_school_id + "/manage_session")

    else:
        session_id = request.POST.get("session_id")
        session_start = request.POST.get("session_start")
        session_end = request.POST.get("session_end")

        try:
            session_year = Session.objects.get(id=session_id)
            session_year.session_start = session_start
            session_year.session_end = session_end
            session_year.save()

            messages.success(request, "Session Year Updated Successfully.")

        except:
            messages.error(request, "Failed to Update Session.")

        finally:
            return redirect("/" + user_school_id + "/manage_session")


def delete_session(request, user_school_id, session_id):
    session = Session.objects.get(id=session_id)
    try:
        session.delete()
        messages.success(request, "Session Deleted Successfully.")

    except:
        messages.error(request, "Failed to Delete Session.")

    finally:
        return redirect("/" + user_school_id + "/manage_session")


# FEES
def manage_fee(request, user_school_id):
    fee_list = Fee.objects.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "fee_list": fee_list,
    }
    return render(request, "admin_templates/manage_fee_template.html", context)


def add_fee(request, user_school_id):
    form = FeeForm()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "form": form,
    }

    return render(request, "admin_templates/add_fee_template.html", context)


def add_fee_save(request, user_school_id):
    if request.method == "POST":
        form = FeeForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Successfully Added Fee")
            # return HttpResponseRedirect(reverse("add_fee"))
            return redirect("/" + user_school_id + "/add_fee")
        else:
            form = FeeForm(request.POST)
            return redirect("/" + user_school_id + "/add_fee")
    else:
        # return HttpResponse("Method Not Allowed")
        messages.error(request, "Invalid Method")
        return redirect("/" + user_school_id + "/add_fee")


def edit_fee(request, user_school_id, fee_id):
    fee = Fee.objects.get(id=fee_id)
    form = FeeForm(request.POST or None, instance=fee)
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "fee": fee,
        "form": form,
    }
    if form.is_valid():
        form.save()
        messages.success(request, "Successfully Edited Fee")
        # return redirect(reverse("manage_fee"))
        return redirect("/" + user_school_id + "/manage_fee")
    return render(request, "admin_templates/edit_fee_template.html", context)


def delete_fee(request, user_school_id, fee_id):
    fee = Fee.objects.get(id=fee_id)
    try:
        fee.delete()
        messages.success(request, "Fee Deleted Successfully.")

    except:
        messages.error(request, "Failed to Delete Fee.")

    finally:
        return redirect("/" + user_school_id + "/manage_fee")


# OTHER VIEWS
def view_fee_payments(request, user_school_id, student_school_id):
    student = User.objects.get(school_id=student_school_id)
    payment_all = Payment.objects.filter(student=student.student, verified=True)
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "payment_all": payment_all,
    }
    return render(request, "admin_templates/student_payment_history.html", context)


def student_records(request, user_school_id, course_id):
    course = ClassLevel.objects.get(class_level_name=course_id)
    students = course.student_set.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "course": course,
        "students": students,
    }
    return render(request, "admin_templates/student_list.html", context)


def student_records_doc(request, course_id):
    # pip install python-docx
    # import docx
    # from django.http import HttpResponse
    # from cStringIO import StringIO

    course = ClassLevel.objects.get(id=course_id)
    students = course.student_set.all()

    doc = docx.Document()
    doc.add_heading("Students in " + str(course.class_level_name), 0)
    # student_list = []
    # students=Student.objects.all().order_by('-course_id.created_at')
    student_list = []
    for student in students:
        small_student = (
            student.user.school_id,
            student.user.first_name,
            student.user.last_name,
        )
        student_list.append(small_student)

    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells
    row[0].text = "SCHOOL ID"
    # row[1].text = 'CLASS'
    row[1].text = "FIRST NAME"
    row[2].text = "LAST NAME"
    # row[3].text = 'LOGIN PASSWORD'
    # print(student_list)

    for id, first, last in student_list:
        row = table.add_row().cells
        row[0].text = str(id)
        # row[1].text = str(course)
        row[1].text = str(first)
        row[2].text = str(last)
        # row[3].text = str(password)

    # f = StringIO(doc)
    # doc.save(f)
    # length = f.tell()
    # f.seek(0)
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename={}.docx".format(
        course.class_level_name
    )
    doc.save(response)

    return response

    # method 2
    # import io
    # buffer = io.BytesIO()
    # doc.save(buffer)
    # buffer.seek(0)
    # response = StreamingHttpResponse(streaming_content=buffer,
    # content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    # response['Content_Disposition'] = 'attachment;filename=download.docx'
    # response['Content_Encoding'] = 'UTF-8'
    # return response


def change_class_level(request, user_school_id, class_level_name):
    if class_level_name == "completed":
        students = Student.objects.filter(student_status=2)
        class_levels = ClassLevel.objects.all()

        context = {
            "user_first_name": request.session.get("user_first_name"),
            "user_last_name": request.session.get("user_last_name"),
            "user_other_names": request.session.get("user_other_names"),
            "user_school_id": user_school_id,
            "class_levels": class_levels,
            "students": students,
            "change_class_level": True,
        }

        return render(request, "admin_templates/students_completed.html", context)

    if class_level_name == "left":
        students = Student.objects.filter(student_status=3)
        class_levels = ClassLevel.objects.all()

        context = {
            "user_first_name": request.session.get("user_first_name"),
            "user_last_name": request.session.get("user_last_name"),
            "user_other_names": request.session.get("user_other_names"),
            "user_school_id": user_school_id,
            "class_levels": class_levels,
            "students": students,
            "change_class_level": True,
        }

        return render(request, "admin_templates/students_left.html", context)

    # student_status_choices = Student.student_status_choices
    class_level = ClassLevel.objects.get(class_level_name=class_level_name)
    class_levels = ClassLevel.objects.all()
    students = class_level.student_set.all()
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "class_level": class_level,
        "class_levels": class_levels,
        "students": students,
        "change_class_level": True,
    }

    return render(request, "admin_templates/manage_students_template.html", context)


def manage_students_completed(request, user_school_id):
    students_completed = Student.objects.filter(class_level=None, student_status=2)
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "students": students_completed,
    }

    return render(request, "admin_templates/students_completed.html", context)


def manage_students_left(request, user_school_id):
    students_left = Student.objects.filter(class_level=None, student_status=3)
    context = {
        "user_first_name": request.session.get("user_first_name"),
        "user_last_name": request.session.get("user_last_name"),
        "user_other_names": request.session.get("user_other_names"),
        "user_school_id": user_school_id,
        "students": students_left,
    }

    return render(request, "admin_templates/students_left.html", context)


def change_class_level_save(request, user_school_id):

    current_session = Session.objects.get(current_session=True)
    if request.method != "POST":
        print("failed")
        return redirect("/" + user_school_id + "/manage_class")

    data = json.loads(request.POST["data"])

    if data.get("class_id") == "completed":

        try:
            with transaction.atomic():
                for student_school_user_id in data.get("students"):
                    new_student = User.objects.get(school_id=student_school_user_id)
                    new_student.student.isOld = False
                    new_student.student.student_status = 2
                    new_student.student.previous_class = None
                    new_student.student.session_completed = str(current_session).split(
                        " "
                    )[0]
                    new_student.student.class_level = None
                    new_student.save()

            messages.success(request, "Student(s) moved successfully")
            return JsonResponse(
                json.dumps(
                    {
                        "redirectUrl": "/"
                        + user_school_id
                        + "/completed"
                        + "/manage_students"
                    }
                ),
                content_type="application/json",
                safe=False,
            )

        except Exception as e:
            print("error: ", e)
            messages.error(request, "Failed to move Student(s)")
            old_class = ClassLevel.objects.get(
                id=User.objects.get(
                    school_id=data.get("students")[0]
                ).student.class_level.id
            )

            return JsonResponse(
                json.dumps(
                    {
                        "/"
                        + user_school_id
                        + "/manage_class/"
                        + old_class.class_level_name
                        + "/manage_students"
                    }
                ),
                content_type="application/json",
                safe=False,
            )

    if data.get("class_id") == "left":

        old_class = ClassLevel.objects.get(
            id=User.objects.get(
                school_id=data.get("students")[0]
            ).student.class_level.id
        )
        try:
            with transaction.atomic():
                for student_school_user_id in data.get("students"):
                    new_student = User.objects.get(school_id=student_school_user_id)
                    new_student.student.isOld = False
                    new_student.student.student_status = 3
                    new_student.student.previous_class = str(old_class)
                    new_student.student.session_completed = None
                    new_student.student.class_level = None
                    new_student.save()

            messages.success(request, "Student(s) moved successfully")
            return JsonResponse(
                json.dumps(
                    {"redirectUrl": "/" + user_school_id + "/left" + "/manage_students"}
                ),
                content_type="application/json",
                safe=False,
            )

        except Exception as e:
            print("error: ", e)
            messages.error(request, "Failed to move Student(s)")

            return JsonResponse(
                json.dumps(
                    {
                        "/"
                        + user_school_id
                        + "/manage_class/"
                        + old_class.class_level_name
                        + "/manage_students"
                    }
                ),
                content_type="application/json",
                safe=False,
            )

    new_class = ClassLevel.objects.get(id=data.get("class_id"))
    old_students = Student.objects.filter(class_level=new_class)
    old_class = User.objects.get(school_id=data.get("students")[0]).student.class_level

    try:
        with transaction.atomic():
            for student in old_students:
                student.isOld = True
                student.save()

            for student_school_user_id in data.get("students"):
                # change old_class id to "" if moving from completed or left to existing class
                new_student = User.objects.get(school_id=student_school_user_id)
                new_student.student.isOld = False
                new_student.student.student_status = 1
                new_student.student.session_completed = None
                new_student.student.previous_class = (
                    str(old_class.id) if old_class else None
                )
                new_student.student.class_level = new_class
                new_student.save()

        messages.success(request, "Student(s) moved successfully")
        return JsonResponse(
            json.dumps(
                {
                    "redirectUrl": "/"
                    + user_school_id
                    + "/manage_class/"
                    + new_class.class_level_name
                    + "/manage_students"
                }
            ),
            content_type="application/json",
            safe=False,
        )

    except Exception as e:
        print("error: ", e)
        messages.error(request, "Failed to move Student(s)")

        # if moving student from existing class to another, run this block
        if User.objects.get(school_id=data.get("students")[0]).student.class_level:

            old_class = ClassLevel.objects.get(
                id=User.objects.get(
                    school_id=data.get("students")[0]
                ).student.class_level.id
            )

            return JsonResponse(
                json.dumps(
                    {
                        "redirectUrl": "/"
                        + user_school_id
                        + "/manage_class/"
                        + old_class.class_level_name
                        + "/manage_students"
                    }
                ),
                content_type="application/json",
                safe=False,
            )

        # if moving student from left or completed, run this block
        if (
            User.objects.get(school_id=data.get("students")[0]).student.student_status
            == 2
        ):

            return JsonResponse(
                json.dumps(
                    {
                        "redirectUrl": "/"
                        + user_school_id
                        + "/completed"
                        + "/manage_students"
                    }
                ),
                content_type="application/json",
                safe=False,
            )

        return JsonResponse(
            json.dumps(
                {"redirectUrl": "/" + user_school_id + "/left" + "/manage_students"}
            ),
            content_type="application/json",
            safe=False,
        )


# def class_search(request, user_school_id):
# if request.method == "POST":
#    searched = request.POST['searched']
#    class_search = ClassLevel.objects.filter(class_level_name__contains=searched)
#    context = {
#    'user_first_name': request.session.get('user_first_name'),
#    'user_last_name': request.session.get('user_last_name'),
#    'user_other_names': request.session.get('user_other_names'),
#    "user_school_id": user_school_id,
#    'searched':searched,
#    'class_search':class_search
# }
#    return render(request,'admin_templates/class_search_template.html', context)
# else:
#    return render(request,"admin_templates/class_search_template.html",{})

# def subject_search(request, user_school_id):
#   if request.method == "POST":
#       searched = request.POST['searched']
#       subject_search = Subject.objects.filter(subject_name__contains=searched)
#       context = {
#       'user_first_name': request.session.get('user_first_name'),
#       'user_last_name': request.session.get('user_last_name'),
#       'user_other_names': request.session.get('user_other_names'),
#       "user_school_id": user_school_id,
#       'searched':searched,
#       'subject_search':subject_search
#   }
#       return render(request,'admin_templates/subject_search_template.html', context)
#   else:
#       return render(request,"admin_templates/subject_search_template.html",{})

# def staff_search(request, user_school_id):
#    if request.method == "POST":
#        searched = request.POST['searched']
#        staff_search = Staff.objects.filter(user.first_name__contains=searched)
#        context = {
#        'user_first_name': request.session.get('user_first_name'),
#        'user_last_name': request.session.get('user_last_name'),
#        'user_other_names': request.session.get('user_other_names'),
#        "user_school_id": user_school_id,
#        'searched':searched,
#        'staff_search':staff_search
#    }
#        return render(request,'admin_templates/staff_search_template.html', context)
#    else:
#        return render(request,"admin_templates/staff_search_template.html",{})

# def student_search(request, user_school_id):
#    if request.method == "POST":
#        searched = request.POST['searched']
#        student_search = Student.objects.filter(user.school_id__contains=searched, class_level.class_level_name__contains=searched)
#        context = {
#        'user_first_name': request.session.get('user_first_name'),
#        'user_last_name': request.session.get('user_last_name'),
#        'user_other_names': request.session.get('user_other_names'),
#        "user_school_id": user_school_id,
#        'searched':searched,
#        'student_search':student_search
#    }
#        return render(request,'admin_templates/student_search_template.html', context)
#    else:
#        return render(request,"admin_templates/student_search_template.html",{})
